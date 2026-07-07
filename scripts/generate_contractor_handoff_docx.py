#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "exports/装修队强弱电必要确认_2026-07-07.docx"
MAIN_IMAGE = ROOT / "diagrams/13_强弱电点位图_施工队简版.png"


FONT = "PingFang SC"
BLUE = RGBColor(31, 79, 120)
MUTED = RGBColor(90, 90, 90)


def set_run_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size: float, color: RGBColor | None = None, bold: bool = False) -> None:
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    set_run_font(run, 15 if level == 1 else 12.5, bold=True, color=BLUE)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, 10.5)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 10.2)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    set_style_font(doc.styles["Normal"], 10.5)
    set_style_font(doc.styles["Heading 1"], 15, BLUE, True)
    set_style_font(doc.styles["Heading 2"], 12.5, BLUE, True)
    set_style_font(doc.styles["List Bullet"], 10.5)
    set_style_font(doc.styles["List Number"], 10.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("强弱电点位必要确认")
    set_run_font(r, 22, bold=True, color=BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    r = sub.add_run("发送对象：设计师 / 装修队 / 电工    日期：2026-07-07")
    set_run_font(r, 10.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(MAIN_IMAGE), width=Cm(26.0))

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(0)
    r = note.add_run("请以图中编号为准；位置为示意，正式落图由设计师按梁位、吊顶、柜体、家具和规范微调。")
    set_run_font(r, 9.5, color=MUTED)

    doc.add_page_break()

    add_heading(doc, "今天只需要确认这些", 1)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    r = intro.add_run("请不要把旧版大包作为今天施工交底依据；本文件只补充设计师当前强弱电图中需要业主额外确认的点。")
    set_run_font(r, 10.5)

    items = [
        "M1 衣帽间机柜：强电插座 + 弱电汇聚。",
        "M2 P2S：独立插座；不接 UPS；不留墙面网口。",
        "M3 AP1/AP2：各 1 根 PoE 网线回衣帽间。",
        "M4 PoE Zigbee：中央高位 1 根网线。",
        "M5 书房桌位：2 网口 + 大容量插座。",
        "M6 次卧/儿童房：2 网口 + 插座。",
        "M7 主卧阳台：至少 1 网口 + 插座。",
        "M8 影音柜：1 网口 + 多位插座。",
        "M9 短焦投影/幕布：不留顶面投影电源；影音柜和幕布盒预留电源/管线。",
        "M10 RS-485：从衣帽间机柜预留到空调/加湿接入点；图中仅示意两端，不代表实际走线。",
        "M11A/M11B：次卫洗衣 + 厨下/总水阀，水电可检修。",
        "M12 窗帘盒：每组电机侧 220V 常电。",
    ]
    for item in items:
        add_number(doc, item)

    add_heading(doc, "本轮明确不要施工", 2)
    for item in [
        "入户高位摄像头 / 低位门铃。",
        "P2S 墙面网口。",
        "AP-R / 第 3 个 AP。",
        "CO2、门磁、水浸等电池传感器逐点拉线。",
        "短焦投影仪顶面吊装电源。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "建议今天发送的材料", 2)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    rows = [
        ("必须发送", "本 DOCX + 13_强弱电点位图_施工队简版.png"),
        ("可选发送", "12_强弱电点位图_业主必要标注.png，只有设计师需要看箭头解释时再发"),
        ("不要发送", "7 月 5 日大包和长篇架构文档，除非对方主动要求背景材料"),
    ]
    for row, (label, detail) in zip(table.rows, rows):
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(20.0)
        set_cell_shading(row.cells[0], "E8EEF5")
        for cell, text, bold in [(row.cells[0], label, True), (row.cells[1], detail, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, 10.2, bold=bold)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
